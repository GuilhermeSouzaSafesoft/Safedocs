from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from styles import PARAGRAPH_STYLES, RUN_STYLES, BLUE_HEX
from utils import hex_to_rgb, set_cell_shading


def apply_run_style(run, style_name="body"):
    style = RUN_STYLES.get(style_name, RUN_STYLES["body"])

    run.font.name = style["font_name"]
    run.font.size = Pt(style["font_size"])
    run.bold = style.get("bold", False)
    run.italic = style.get("italic", False)
    run.font.color.rgb = hex_to_rgb(style["color"])


def apply_paragraph_format(paragraph, style_name="body"):
    style = PARAGRAPH_STYLES.get(style_name, PARAGRAPH_STYLES["body"])
    align = style.get("align", "left")

    if align == "center":
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif align == "right":
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    elif align == "justify":
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    else:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    pf = paragraph.paragraph_format
    pf.space_before = Pt(style.get("space_before", 0))
    pf.space_after = Pt(style.get("space_after", 0))


def add_text_with_style(paragraph, text, style_name="body"):
    run = paragraph.add_run(text)
    apply_run_style(run, style_name)
    return run


def render_paragraph(doc, block):
    p = doc.add_paragraph()

    paragraph_style_name = block.get("style", "body")
    apply_paragraph_format(p, paragraph_style_name)

    if "runs" in block:
        for item in block["runs"]:
            text = item.get("text", "")
            run_style_name = item.get("style", "body")
            add_text_with_style(p, text, run_style_name)
    else:
        text = block.get("text", "")
        run_style_name = block.get("run_style", paragraph_style_name)

        if run_style_name not in RUN_STYLES:
            run_style_name = "body"

        add_text_with_style(p, text, run_style_name)


def render_heading(doc, block):
    level = int(block.get("level", 1))
    text = block.get("text", "")

    if level <= 1:
        style_key = "heading_1"
    elif level == 2:
        style_key = "heading_2"
    elif level == 3:
        style_key = "heading_3"
    else:
        style_key = "heading_4"

    p = doc.add_paragraph()
    apply_paragraph_format(p, style_key)
    add_text_with_style(p, text, style_key)


def render_bullets(doc, block):
    items = block.get("items", [])

    for item in items:
        p = doc.add_paragraph()
        apply_paragraph_format(p, "bullet")
        add_text_with_style(p, "• ", "bold")
        add_text_with_style(p, str(item), "bullet")


def render_numbered(doc, block):
    items = block.get("items", [])

    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        apply_paragraph_format(p, "number")
        add_text_with_style(p, f"{index}. ", "bold")
        add_text_with_style(p, str(item), "number")


def render_table(doc, block):
    cols = block.get("columns", [])
    rows = block.get("rows", [])

    if not cols:
        p = doc.add_paragraph()
        apply_paragraph_format(p, "body")
        add_text_with_style(p, "[Tabela ignorada: sem colunas]", "body")
        return

    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for col_index, col_name in enumerate(cols):
        cell = header_cells[col_index]
        set_cell_shading(cell, BLUE_HEX)

        p = cell.paragraphs[0]
        apply_paragraph_format(p, "table_header")
        add_text_with_style(p, str(col_name), "table_header")

    for row_data in rows:
        row_cells = table.add_row().cells
        for col_index in range(len(cols)):
            value = str(row_data[col_index]) if col_index < len(row_data) else ""

            p = row_cells[col_index].paragraphs[0]
            apply_paragraph_format(p, "table_cell")
            add_text_with_style(p, value, "table_cell")


def render_image_caption(doc, block):
    p = doc.add_paragraph()
    apply_paragraph_format(p, "caption")
    add_text_with_style(p, block.get("text", ""), "caption")


def render_section_divider(doc):
    p = doc.add_paragraph()
    apply_paragraph_format(p, "body")
    add_text_with_style(p, "─" * 50, "blue")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def render_block(doc, block):
    block_type = block.get("type")

    if block_type == "title":
        title_block = dict(block)
        title_block["style"] = "title"
        title_block["run_style"] = "title"
        render_paragraph(doc, title_block)

    elif block_type == "paragraph":
        render_paragraph(doc, block)

    elif block_type == "heading":
        render_heading(doc, block)

    elif block_type == "bullets":
        render_bullets(doc, block)

    elif block_type == "numbered":
        render_numbered(doc, block)

    elif block_type == "table":
        render_table(doc, block)

    elif block_type == "image_caption":
        render_image_caption(doc, block)

    elif block_type == "page_break":
        doc.add_page_break()

    else:
        p = doc.add_paragraph()
        apply_paragraph_format(p, "body")
        add_text_with_style(p, f"[Bloco não suportado: {block_type}]", "body")

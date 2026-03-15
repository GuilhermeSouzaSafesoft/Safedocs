from docx import Document

from blocks import render_block, apply_paragraph_format, add_text_with_style
from pathlib import Path
from docx.shared import Cm

COVER_SPACER_LINES = 8
COVER_IMAGE_WIDTH_CM = 10

def validar_arquivos(template_file, json_file):
    if not template_file.exists():
        raise FileNotFoundError(f"Template não encontrado: '{template_file}'")

    if not json_file.exists():
        raise FileNotFoundError(f"JSON não encontrado: '{json_file}'")


def limpar_paragrafos(doc):
    while doc.paragraphs:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)


def add_cover_line(doc, text, paragraph_style, run_style):
    p = doc.add_paragraph()
    apply_paragraph_format(p, paragraph_style)
    add_text_with_style(p, text, run_style)
    return p


def add_cover_spacing(doc, lines=COVER_SPACER_LINES):
    for _ in range(lines):
        doc.add_paragraph()

def add_cover_image(doc, image_path, width_cm=COVER_IMAGE_WIDTH_CM):
    if not image_path or not Path(image_path).exists():
        raise FileNotFoundError(f"Imagem da capa não encontrada: '{image_path}'")

    p = doc.add_paragraph()
    apply_paragraph_format(p, "image")

    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))


def gerar_capa(doc, meta, placeholder_image=None):
    document_type = meta.get("document_type", "").upper()
    item_name = meta.get("item_name", "")
    document_code = meta.get("document_code", "")
    revision = meta.get("revision", "")
    date = meta.get("date", "")

    cover_image_path = meta.get("cover_image_path") or placeholder_image

    add_cover_line(doc, document_type, "cover_document_type", "cover_document_type")
    add_cover_line(doc, item_name, "cover_item_name", "cover_item_name")

    add_cover_spacing(doc, lines=2)
    add_cover_image(doc, cover_image_path)
    add_cover_spacing(doc, lines=2)

    add_cover_line(doc, f"Código: {document_code}", "cover_meta", "cover_meta")
    add_cover_line(doc, f"Data: {date}", "cover_meta", "cover_meta")
    add_cover_line(doc, f"Versão: {revision}", "cover_meta", "cover_meta")

    doc.add_page_break()

def gerar_documento(dados_json, template_file, output_file, placeholder_image=None):
    doc = Document(str(template_file))
    limpar_paragrafos(doc)

    gerar_capa(doc, dados_json["meta"], placeholder_image=placeholder_image)

    for block in dados_json["blocks"]:
        render_block(doc, block, placeholder_image=placeholder_image)

    doc.save(str(output_file))
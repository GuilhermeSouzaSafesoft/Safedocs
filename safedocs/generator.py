from __future__ import annotations

from pathlib import Path
from typing import Union

from docx import Document

from .blocks import render_block
from .cover import gerar_capa
from .models import DocumentData


def validar_arquivos(template_file: Path, json_file: Path) -> None:
    if not template_file.exists():
        raise FileNotFoundError(f"Template não encontrado: '{template_file}'")

    if not json_file.exists():
        raise FileNotFoundError(f"JSON não encontrado: '{json_file}'")


def limpar_paragrafos(doc: Document) -> None:
    while doc.paragraphs:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)


def gerar_documento(
    dados: DocumentData,
    template_file: Union[str, Path],
    output_file: Union[str, Path],
    placeholder_image: Union[str, Path, None] = None,
) -> None:
    """
    Gera um documento DOCX a partir dos dados tipados e de um template.
    """
    doc = Document(str(template_file))
    limpar_paragrafos(doc)

    gerar_capa(doc, dados.meta, placeholder_image=placeholder_image)

    for block in dados.blocks:
        render_block(doc, block.raw, placeholder_image=placeholder_image)

    doc.save(str(output_file))
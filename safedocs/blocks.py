from typing import Any, Dict, Iterable, List, Mapping, Sequence

from pathlib import Path

from docx.document import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from .styles import BLUE_HEX, PARAGRAPH_STYLES, RUN_STYLES
from .utils import hex_to_rgb, set_cell_shading


def apply_run_style(run: Run, style_name: str = "body") -> None:
    style = RUN_STYLES.get(style_name, RUN_STYLES["body"])

    run.font.name = style["font_name"]
    run.font.size = Pt(style["font_size"])
    run.bold = style.get("bold", False)
    run.italic = style.get("italic", False)
    run.font.color.rgb = hex_to_rgb(style["color"])


def apply_paragraph_format(paragraph: Paragraph, style_name: str = "body") -> None:
    style = PARAGRAPH_STYLES.get(style_name, PARAGRAPH_STYLES["body"])
    align = style.get("align", "left")

    if align == "center":
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif align == "right":
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    elif align == "justify":
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    else:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    pf = paragraph.paragraph_format
    pf.space_before = Pt(style.get("space_before", 0))
    pf.space_after = Pt(style.get("space_after", 0))


def add_text_with_style(paragraph: Paragraph, text: str, style_name: str = "body") -> Run:
    run = paragraph.add_run(text)
    apply_run_style(run, style_name)
    return run


def render_paragraph(doc: Document, block: Mapping[str, Any]) -> None:
    p = doc.add_paragraph()

    paragraph_style_name = block.get("style", "body")
    apply_paragraph_format(p, paragraph_style_name)

    runs = block.get("runs")
    if runs:
        for item in runs:
            text = item.get("text", "")
            run_style_name = item.get("style", "body")
            add_text_with_style(p, text, run_style_name)
    else:
        text = block.get("text", "")
        run_style_name = block.get("run_style", paragraph_style_name)

        if run_style_name not in RUN_STYLES:
            run_style_name = "body"

        add_text_with_style(p, text, run_style_name)


def render_heading_block(doc: Document, block: Mapping[str, Any], style_key: str) -> None:
    text = block.get("text", "")

    p = doc.add_paragraph()
    apply_paragraph_format(p, style_key)
    add_text_with_style(p, text, style_key)


def render_bullets(doc: Document, block: Mapping[str, Any]) -> None:
    items: Iterable[Any] = block.get("items") or []

    for item in items:
        p = doc.add_paragraph()
        apply_paragraph_format(p, "bullet")
        add_text_with_style(p, "• ", "bold")
        add_text_with_style(p, str(item), "bullet")


def render_numbered(doc: Document, block: Mapping[str, Any]) -> None:
    items: Sequence[Any] = block.get("items") or []

    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        apply_paragraph_format(p, "number")
        add_text_with_style(p, f"{index}. ", "bold")
        add_text_with_style(p, str(item), "number")


def render_table(doc: Document, block: Mapping[str, Any]) -> None:
    cols: List[str] = list(block.get("columns") or [])
    rows: List[Sequence[Any]] = list(block.get("rows") or [])

    if not cols:
        p = doc.add_paragraph()
        apply_paragraph_format(p, "body")
        add_text_with_style(p, "[Tabela ignorada: sem colunas]", "body")
        return

    table: Table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for col_index, col_name in enumerate(cols):
        cell = header_cells[col_index]
        set_cell_shading(cell, BLUE_HEX)

        p = cell.paragraphs[0]
        apply_paragraph_format(p, "table_header")
        add_text_with_style(p, str(col_name), "table_header")

    for row_data in rows:
        row_cells = table.add_row().cells
        for col_index in range(len(cols)):
            value = str(row_data[col_index]) if col_index < len(row_data) else ""

            p = row_cells[col_index].paragraphs[0]
            apply_paragraph_format(p, "table_cell")
            add_text_with_style(p, value, "table_cell")


def render_image_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    apply_paragraph_format(p, "caption")
    add_text_with_style(p, text, "caption")


def render_image(doc: Document, block: Mapping[str, Any], placeholder_image: Path | str | None = None) -> None:
    image_path = str(block.get("path", "") or "")
    # Largura fixa para imagens de conteúdo
    width_cm: float = 8.0
    caption = str(block.get("caption", "") or "")

    final_image_path: str | None = None

    if image_path and Path(image_path).exists():
        final_image_path = image_path
    elif placeholder_image and Path(str(placeholder_image)).exists():
        final_image_path = str(placeholder_image)

    if final_image_path is None:
        p = doc.add_paragraph()
        apply_paragraph_format(p, "body")
        add_text_with_style(p, "[Imagem não encontrada]", "body")

        if caption:
            render_image_caption(doc, caption)
        return

    p = doc.add_paragraph()
    apply_paragraph_format(p, "image")

    run = p.add_run()
    run.add_picture(final_image_path, width=Cm(width_cm))

    if caption:
        render_image_caption(doc, caption)



def render_block(doc: Document, block: Dict[str, Any], placeholder_image: Path | str | None = None) -> None:
    block_type = block.get("type")

    if block_type == "paragrafo":
        render_paragraph(doc, block)
    elif block_type == "secao":
        render_heading_block(doc, block, "secao")
    elif block_type == "subsecao":
        render_heading_block(doc, block, "subsecao")
    elif block_type == "subsubsecao":
        render_heading_block(doc, block, "subsubsecao")
    elif block_type == "subsubsubsecao":
        render_heading_block(doc, block, "subsubsubsecao")
    elif block_type == "lista_com_marcadores":
        render_bullets(doc, block)
    elif block_type == "lista_numerada":
        render_numbered(doc, block)
    elif block_type == "tabela":
        render_table(doc, block)
    elif block_type == "imagem":
        render_image(doc, block, placeholder_image=placeholder_image)
    elif block_type == "quebra_de_pagina":
        doc.add_page_break()
    else:
        p = doc.add_paragraph()
        apply_paragraph_format(p, "body")
        add_text_with_style(p, f"[Bloco não suportado: {block_type}]", "body")

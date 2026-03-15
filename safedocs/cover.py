from __future__ import annotations

from pathlib import Path
from typing import Optional, Union

from docx import Document
from docx.shared import Cm

from .blocks import add_text_with_style, apply_paragraph_format
from .models import Meta

COVER_SPACER_LINES = 8
COVER_IMAGE_WIDTH_CM = 10


def _add_cover_line(doc: Document, text: str, paragraph_style: str, run_style: str):
    p = doc.add_paragraph()
    apply_paragraph_format(p, paragraph_style)
    add_text_with_style(p, text, run_style)
    return p


def _add_cover_spacing(doc: Document, lines: int = COVER_SPACER_LINES) -> None:
    for _ in range(lines):
        doc.add_paragraph()


def _add_cover_image(
    doc: Document,
    image_path: Union[str, Path],
    width_cm: float = COVER_IMAGE_WIDTH_CM,
) -> None:
    p = doc.add_paragraph()
    apply_paragraph_format(p, "image")

    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))


def gerar_capa(
    doc: Document,
    meta: Meta,
    placeholder_image: Optional[Union[str, Path]] = None,
) -> None:
    """
    Gera a página de capa do documento a partir dos metadados.
    """
    document_type = meta.tipo_do_documento.upper()
    item_name = meta.nome_do_item
    document_code = meta.codigo
    revision = meta.revisao
    date = meta.data

    placeholder_path: Optional[Union[str, Path]] = None
    if placeholder_image is not None and Path(str(placeholder_image)).exists():
        placeholder_path = placeholder_image

    _add_cover_line(doc, document_type, "cover_document_type", "cover_document_type")
    _add_cover_line(doc, item_name, "cover_item_name", "cover_item_name")

    _add_cover_spacing(doc, lines=2)
    if placeholder_path is not None:
        _add_cover_image(doc, placeholder_path)
    _add_cover_spacing(doc, lines=2)

    _add_cover_line(doc, f"Código: {document_code}", "cover_meta", "cover_meta")
    _add_cover_line(doc, f"Data: {date}", "cover_meta", "cover_meta")
    _add_cover_line(doc, f"Versão: {revision}", "cover_meta", "cover_meta")

    doc.add_page_break()

